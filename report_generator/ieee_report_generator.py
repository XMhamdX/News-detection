import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class IEEEReportGenerator:
    def __init__(self):
        self.document = Document()
        self.setup_document_format()

    def setup_document_format(self):
        # Set margins
        section = self.document.sections[0]
        section.top_margin = Cm(1.78)  # 0.7 inches
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.78)  # 0.7 inches
        section.right_margin = Cm(1.78)  # 0.7 inches
        
        # Set to two columns
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')  # space between columns (0.5 inch = 720 twips)

        # Set up the default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)

    def add_title_block(self, title):
        # Add title - spans both columns
        section = self.document.add_section(WD_SECTION.CONTINUOUS)
        section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')  # Single column for title
        
        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.name = 'Times New Roman'
        title_run.bold = True

        # Return to two columns
        section = self.document.add_section(WD_SECTION.CONTINUOUS)
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')

    def add_authors(self, authors_data):
        authors_para = self.document.add_paragraph()
        authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, author in enumerate(authors_data):
            if i > 0:
                authors_para.add_run('\n')
            
            # Add author name
            name_run = authors_para.add_run(f"line {i+1}: {author['name']}\n")
            name_run.font.size = Pt(10)
            
            # Add department
            dept_run = authors_para.add_run(f"line {i+2}: {author['dept']}\n")
            dept_run.italic = True
            dept_run.font.size = Pt(8)
            
            # Add organization
            org_run = authors_para.add_run(f"line {i+3}: {author['org']}\n")
            org_run.font.size = Pt(8)
            
            # Add location
            loc_run = authors_para.add_run(f"line {i+4}: {author['location']}\n")
            loc_run.font.size = Pt(8)
            
            # Add email/ORCID
            email_run = authors_para.add_run(f"line {i+5}: {author['email']}")
            email_run.font.size = Pt(8)

    def add_abstract(self, abstract_text, keywords):
        abstract_para = self.document.add_paragraph()
        
        # Add Abstract heading
        abstract_heading = abstract_para.add_run('Abstract—')
        abstract_heading.bold = True
        abstract_heading.font.size = Pt(10)
        
        # Add abstract text
        abstract_content = abstract_para.add_run(abstract_text)
        abstract_content.font.size = Pt(9)
        
        # Add Keywords
        keywords_para = self.document.add_paragraph()
        keywords_run = keywords_para.add_run('Keywords—')
        keywords_run.italic = True
        keywords_run.font.size = Pt(9)
        keywords_para.add_run(keywords)
        keywords_para.add_run('\n\n')

    def add_section(self, number, title, content):
        # Add section heading with Roman numeral
        heading = self.document.add_paragraph()
        heading_run = heading.add_run(f"{number}. {title}")
        heading_run.bold = True
        heading_run.font.size = Pt(10)

        # Add section content
        content_para = self.document.add_paragraph()
        content_run = content_para.add_run(content)
        content_run.font.size = Pt(10)

    def add_equation(self, equation, number):
        eq_para = self.document.add_paragraph()
        eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_run = eq_para.add_run(equation)
        eq_run.font.italic = True
        
        # Add equation number in parentheses
        eq_para.add_run(f" ({number})")

    def save(self, output_path):
        self.document.save(output_path)

def generate_sample_report():
    report = IEEEReportGenerator()

    # Title
    report.add_title_block("Derin Öğrenme Kullanarak Haber Sınıflandırma ve Analizi")

    # Authors
    authors_data = [
        {
            'name': 'Birinci Yazar',
            'dept': 'Bilgisayar Mühendisliği Bölümü',
            'org': 'Üniversite Adı',
            'location': 'Şehir, Ülke',
            'email': 'email@university.edu.tr'
        },
        {
            'name': 'İkinci Yazar',
            'dept': 'Bilgisayar Mühendisliği Bölümü',
            'org': 'Üniversite Adı',
            'location': 'Şehir, Ülke',
            'email': 'email2@university.edu.tr'
        }
    ]
    report.add_authors(authors_data)

    # Abstract and Keywords
    abstract = """Bu çalışmada, haber makalelerinin otomatik sınıflandırılması için 
    derin öğrenme tabanlı bir yaklaşım sunulmaktadır. Önerilen sistem, doğal dil 
    işleme teknikleri ve modern derin öğrenme modellerini kullanarak haber 
    içeriklerini analiz etmekte ve sınıflandırmaktadır."""
    
    keywords = "derin öğrenme, haber sınıflandırma, doğal dil işleme, BERT"
    report.add_abstract(abstract, keywords)

    # Sections
    sections = [
        ("I", "GİRİŞ", """Günümüzde dijital medyanın hızlı büyümesi, haber içeriklerinin 
        otomatik sınıflandırılması ihtiyacını doğurmuştur."""),
        
        ("II", "YÖNTEM", """A. Veri Hazırlama\nVeri seti, çeşitli haber kaynaklarından 
        toplanan makalelerden oluşmaktadır.\n\nB. Model Mimarisi\nBERT tabanlı bir 
        model kullanılmıştır."""),
        
        ("III", "DENEYSEL SONUÇLAR", """Deneysel sonuçlar, sistemimizin %90'ın 
        üzerinde bir doğruluk oranıyla çalıştığını göstermektedir."""),
        
        ("IV", "SONUÇ", """Bu çalışma, derin öğrenme tekniklerinin haber sınıflandırma 
        alanında etkili bir şekilde kullanılabileceğini göstermiştir.""")
    ]

    for section in sections:
        report.add_section(*section)

    # Add sample equation
    report.add_equation("y = Wx + b", 1)

    # Save the report
    output_path = "IEEE_Report_News_Classification.docx"
    report.save(output_path)
    print(f"Report saved as: {output_path}")

if __name__ == "__main__":
    generate_sample_report()
